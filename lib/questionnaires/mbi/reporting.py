"""
lib/questionnaires/mbi/reporting.py  — rapport Word MBI.
Structure : page de garde / 1. Population / 2. Risque burnout composite /
            3. Résultats par dimension (tableau) / 4. Visualisations / note
"""
import io
from datetime import datetime
from typing import Dict, Any, Optional

from .config import DIMENSIONS, THRESHOLDS

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def _rgb(r,g,b): return RGBColor(r,g,b) if DOCX_AVAILABLE else None
C_DARK=_rgb(15,35,64); C_BLUE=_rgb(56,163,232); C_MID=_rgb(47,87,127)
C_GRAY=_rgb(107,136,168); C_GREEN=_rgb(34,197,94); C_RED=_rgb(239,68,68)
C_ORG=_rgb(249,115,22); C_WHITE=_rgb(255,255,255)

def _cell_bg(cell,hx):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx.lstrip("#")); tcPr.append(shd)

def _cell_borders(cell,color="D6E8F7"):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcB=OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el=OxmlElement(f"w:{side}"); el.set(qn("w:val"),"single")
        el.set(qn("w:sz"),"4"); el.set(qn("w:color"),color.lstrip("#")); tcB.append(el)
    tcPr.append(tcB)

def _spacing(para,before=0,after=80):
    pPr=para._p.get_or_add_pPr(); sp=OxmlElement("w:spacing")
    sp.set(qn("w:before"),str(before)); sp.set(qn("w:after"),str(after)); pPr.append(sp)

def _hex_rgb(h):
    h=h.lstrip("#"); return _rgb(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))

RISK_HEX={"Faible":"22C55E","Modéré":"F97316","Élevé":"EF4444"}
CAT_HEX={"Bas":"22C55E","Modéré":"F97316","Élevé":"EF4444"}


class MBIReporting:
    def __init__(self, company_name="l'organisation"):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx requis.")
        self.company_name=company_name

    def _h(self, doc, text, level=1):
        h=doc.add_heading(text,level=level)
        for run in h.runs: run.font.color.rgb=C_DARK if level==1 else C_MID
        return h

    def _divider(self, doc, color="D6E8F7", thick=6):
        p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
        pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),str(thick))
        bot.set(qn("w:color"),color.lstrip("#")); pBdr.append(bot); pPr.append(pBdr)
        _spacing(p,after=120)

    def _kpi(self, doc, label, pct, n, good=None, inv=False):
        p=doc.add_paragraph(style="Normal"); _spacing(p,before=40,after=40)
        b=p.add_run("• "); b.font.color.rgb=C_BLUE
        l=p.add_run(f"{label} : "); l.font.bold=True; l.font.color.rgb=C_MID
        vc=(C_GREEN if ((pct>=good)!=inv) else C_RED) if good is not None else C_BLUE
        v=p.add_run(f"{pct:.1f}%  (n={n})"); v.font.bold=True; v.font.color.rgb=vc

    def _note(self, doc, text):
        p=doc.add_paragraph(style="Normal"); _spacing(p,before=40,after=40)
        r=p.add_run(text); r.font.italic=True; r.font.color.rgb=C_GRAY; r.font.size=Pt(9)

    def _cover(self, doc, total):
        title=doc.add_heading("Rapport MBI — Burnout",0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.color.rgb=C_DARK; title.runs[0].font.size=Pt(26)
        sub=doc.add_heading("Maslach Burnout Inventory — General Survey (MBI-GS)",level=2)
        sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].font.color.rgb=_rgb(249,115,22); sub.runs[0].font.size=Pt(13)
        doc.add_paragraph(); self._divider(doc,"F97316",thick=10); doc.add_paragraph()
        for key,val in [
            ("Organisation",    self.company_name),
            ("Date",            datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Effectif analysé",f"{total} répondant(s)"),
            ("Référence",       "Schaufeli, Leiter, Maslach & Jackson (1996)"),
            ("Échelle",         "Fréquence 0 (jamais) → 6 (tous les jours)"),
        ]:
            p=doc.add_paragraph(style="Normal"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p,before=40,after=40)
            r1=p.add_run(f"{key} : "); r1.font.bold=True; r1.font.color.rgb=C_MID
            r2=p.add_run(val); r2.font.color.rgb=C_DARK
        doc.add_paragraph(); self._divider(doc,"F97316",thick=10); doc.add_page_break()

    def _s1_pop(self, doc, metrics):
        self._h(doc,"1. Population")
        demo=metrics.get("demographics",{})
        men=demo.get("men",{}); women=demo.get("women",{})
        tbl=doc.add_table(rows=2,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
        headers=["Hommes","Femmes","Âge moyen"]
        values=[f"{men.get('pct',0):.1f}%  (n={men.get('n',0)})",
                f"{women.get('pct',0):.1f}%  (n={women.get('n',0)})",
                f"{demo.get('avg_age',0)} ans"]
        for i,(h,v) in enumerate(zip(headers,values)):
            ch=tbl.rows[0].cells[i]; cv=tbl.rows[1].cells[i]
            ch.width=Inches(2.0); cv.width=Inches(2.0)
            _cell_bg(ch,"C2410C"); _cell_borders(ch,"C2410C")
            _cell_bg(cv,"FEF3C7"); _cell_borders(cv,"FCD34D")
            ph=ch.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rh=ph.add_run(h); rh.font.bold=True; rh.font.color.rgb=C_WHITE; rh.font.size=Pt(10)
            pv=cv.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rv=pv.add_run(v); rv.font.bold=True; rv.font.color.rgb=C_DARK; rv.font.size=Pt(11)
        doc.add_paragraph()

    def _s2_risk(self, doc, metrics):
        self._h(doc,"2. Risque Global de Burnout")
        risk=metrics.get("burnout_risk",{})
        if risk:
            # Tableau 3 cases colorées
            tbl=doc.add_table(rows=2,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
            for i,(level,data) in enumerate(risk.items()):
                hx=RISK_HEX.get(level,"94A3B8")
                ch=tbl.rows[0].cells[i]; cv=tbl.rows[1].cells[i]
                ch.width=Inches(2.0); cv.width=Inches(2.0)
                _cell_bg(ch,hx); _cell_borders(ch,hx)
                ph=ch.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER
                rh=ph.add_run(f"Risque {level}"); rh.font.bold=True
                rh.font.color.rgb=C_WHITE; rh.font.size=Pt(10)
                _cell_bg(cv,"F9FAFB"); _cell_borders(cv,"D6E8F7")
                pv=cv.paragraphs[0]; pv.alignment=WD_ALIGN_PARAGRAPH.CENTER
                rv=pv.add_run(f"{data.get('pct',0):.1f}%")
                rv.font.bold=True; rv.font.size=Pt(16); rv.font.color.rgb=_hex_rgb(hx)
                pv.add_run(f"\nn = {data.get('n',0)}").font.color.rgb=C_MID
        else:
            self._note(doc,"Score composite non disponible — vérifiez les 3 dimensions.")
        self._note(doc,
            "Un profil à risque élevé cumule : épuisement élevé + cynisme élevé + efficacité faible. "
            "Ces seuils sont indicatifs, non diagnostiques.")
        doc.add_paragraph()

    def _s3_dims(self, doc, metrics):
        self._h(doc,"3. Résultats par Dimension")
        p_intro=doc.add_paragraph(style="Normal"); _spacing(p_intro,after=100)
        p_intro.add_run("Note : ").font.bold=True
        p_intro.add_run(
            "Pour l'Épuisement et le Cynisme, un score élevé = risque. "
            "Pour l'Efficacité, un score faible = risque."
        ).font.italic=True

        dims=metrics.get("dimensions",{})
        if not dims: return

        tbl=doc.add_table(rows=1,cols=5); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style="Table Grid"
        for cell,txt in zip(tbl.rows[0].cells,
                             ["Dimension","Score moy.","% Bas","% Modéré","% Élevé"]):
            _cell_bg(cell,"0F2340"); _cell_borders(cell,"0F2340")
            ph=cell.paragraphs[0]; ph.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=ph.add_run(txt); r.font.bold=True; r.font.color.rgb=C_WHITE; r.font.size=Pt(9)

        for i,(dim_key,data) in enumerate(dims.items()):
            row=tbl.add_row(); bg="FFF7ED" if i%2==0 else "FFFFFF"
            cats=data.get("categories",{})
            pct_eleve=cats.get("Élevé",{}).get("pct",0)
            # Pour efficacy : risque = Bas
            is_risk_col = "EF4444" if (dim_key!="efficacy" and pct_eleve>30) else "22C55E"
            vals_row=[
                data["label"],
                f"{data.get('mean',0):.2f}",
                f"{cats.get('Bas',{}).get('pct',0):.1f}%",
                f"{cats.get('Modéré',{}).get('pct',0):.1f}%",
                f"{pct_eleve:.1f}%",
            ]
            for j,(cell,val) in enumerate(zip(row.cells,vals_row)):
                _cell_bg(cell,bg); _cell_borders(cell,"D6E8F7")
                p=cell.paragraphs[0]
                p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
                r=p.add_run(val)
                if j==4: r.font.bold=True; r.font.color.rgb=_hex_rgb(is_risk_col)
                elif j==2: r.font.color.rgb=C_GREEN
                else: r.font.color.rgb=C_DARK
        doc.add_paragraph()

    def _s4_figures(self, doc, figures):
        if not figures: return
        self._h(doc,"4. Visualisations")
        fig_meta=[
            ("burnout_donut","4.1 Risque burnout composite",  Inches(4.5)),
            ("mbi_radar",    "4.2 Radar MBI-GS",              Inches(4.8)),
            ("mbi_bars",     "4.3 Répartition par dimension", Inches(5.5)),
            ("mbi_heatmap",  "4.4 Scores moyens par Direction",Inches(5.2)),
            ("age_pyramid",  "4.5 Pyramide des âges",         Inches(4.8)),
        ]
        for key,title,width in fig_meta:
            if key not in figures or not figures[key]: continue
            self._h(doc,title,level=2)
            doc.add_picture(io.BytesIO(figures[key]),width=width)
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    def generate(self, metrics:Dict[str,Any], figures:Optional[Dict[str,bytes]]=None) -> bytes:
        doc=Document()
        doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
        for s in doc.sections:
            s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5)
            s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)
        total=metrics.get("demographics",{}).get("total","N/A")
        self._cover(doc,total)
        self._s1_pop(doc,metrics)
        self._s2_risk(doc,metrics)
        self._s3_dims(doc,metrics)
        self._s4_figures(doc,figures or {})
        doc.add_paragraph(); self._divider(doc,"D6E8F7",thick=4)
        for note in [
            "Seuils MBI-GS : Épuisement (≥3.0 élevé), Cynisme (≥2.0 élevé), Efficacité (≤3.0 faible). Indicatifs, non cliniques.",
            "Référence : Schaufeli, Leiter, Maslach & Jackson (1996). MBI — General Survey.",
            f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Wave-CI.",
        ]: self._note(doc,f"• {note}")
        buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()
