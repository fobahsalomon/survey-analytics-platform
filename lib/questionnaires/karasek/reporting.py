"""
lib/questionnaires/karasek/reporting.py
Rapport Word Karasek — structure fusionnée des deux versions.

Structure :
  Page de garde
  1. Données Démographiques & Modes de Vie
     1.1 Données démographiques (tableau)
     1.2 Indicateurs de modes de vie
  2. Résultats Principaux
     2.1 Méthodologie & Seuils Théoriques (tableau coloré)
     2.2 Quadrants de Karasek (tableau 4 cases colorées)
     2.3 Prévalence des RPS
     2.4 Indicateurs clés de stress
     2.5 Climat organisationnel (tableau RH coloré)
  3. Visualisations (toutes les figures PNG)
  Note méthodologique
"""

import io
from datetime import datetime
from typing import Dict, Any, Optional

from .config import THRESHOLDS, SCORE_LABELS

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ─── Palette Wave-CI ─────────────────────────────────────────────────────────
def _rgb(r, g, b):
    return RGBColor(r, g, b) if DOCX_AVAILABLE else None

C_DARK   = _rgb(15,  35,  64)
C_BLUE   = _rgb(56,  163, 232)
C_MID    = _rgb(47,  87,  127)
C_GRAY   = _rgb(107, 136, 168)
C_GREEN  = _rgb(34,  197, 94)
C_RED    = _rgb(239, 68,  68)
C_ORANGE = _rgb(249, 115, 22)
C_WHITE  = _rgb(255, 255, 255)

QUAD_HEX = {
    "Tendu":   "EF4444",
    "Actif":   "22C55E",
    "Passif":  "94A3B8",
    "Détendu": "38A3E8",
}


# =============================================================================
# XML HELPERS
# =============================================================================

def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _cell_borders(cell, color: str = "D6E8F7"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color.lstrip("#"))
        tcB.append(el)
    tcPr.append(tcB)


def _spacing(para, before: int = 0, after: int = 80):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"),  str(after))
    pPr.append(sp)


def _hex_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return _rgb(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# =============================================================================
# CLASSE PRINCIPALE
# =============================================================================

class KarasekReporting:
    """Génère le rapport Word Karasek complet — structure fusionnée."""

    def __init__(self, company_name: str = "l'organisation"):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx requis : pip install python-docx --break-system-packages"
            )
        self.company_name = company_name

    # ── micro-helpers ─────────────────────────────────────────────────────────

    def _h(self, doc, text: str, level: int = 1):
        h = doc.add_heading(text, level=level)
        color = C_DARK if level == 1 else C_MID
        for run in h.runs:
            run.font.color.rgb = color
        return h

    def _divider(self, doc, color: str = "D6E8F7", thick: int = 6):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(thick))
        bot.set(qn("w:color"), color.lstrip("#"))
        pBdr.append(bot)
        pPr.append(pBdr)
        _spacing(p, after=120)

    def _kpi(self, doc, label: str, pct: float, n: int,
             good: Optional[float] = None, inv: bool = False,
             color: RGBColor = None):
        p = doc.add_paragraph(style="Normal")
        _spacing(p, before=40, after=40)
        b = p.add_run("• ");     b.font.color.rgb = C_BLUE
        l = p.add_run(f"{label} : ")
        l.font.bold = True;      l.font.color.rgb = C_MID
        if color:
            val_c = color
        elif good is not None:
            val_c = C_GREEN if ((pct >= good) != inv) else C_RED
        else:
            val_c = C_BLUE
        v = p.add_run(f"{pct:.1f}%  (n={n})")
        v.font.bold = True;      v.font.color.rgb = val_c

    def _note(self, doc, text: str):
        p = doc.add_paragraph(style="Normal")
        _spacing(p, before=40, after=40)
        r = p.add_run(text)
        r.font.italic = True; r.font.color.rgb = C_GRAY; r.font.size = Pt(9)

    # ── page de garde ─────────────────────────────────────────────────────────

    def _cover(self, doc, total):
        title = doc.add_heading("Rapport d'Analyse Karasek", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.color.rgb = C_DARK
            title.runs[0].font.size = Pt(26)

        sub = doc.add_heading("Risques Psychosociaux au Travail", level=2)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].font.color.rgb = C_BLUE
            sub.runs[0].font.size = Pt(15)

        doc.add_paragraph()
        self._divider(doc, "38A3E8", thick=10)
        doc.add_paragraph()

        for key, val in [
            ("Organisation",      self.company_name),
            ("Date",              datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Effectif analysé",  f"{total} répondant(s)"),
            ("Seuils",            "Théoriques — point médian Likert 1–4 (non cliniques)"),
            ("Modèle",            "Karasek & Theorell — Demande–Contrôle–Soutien (DCS)"),
        ]:
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacing(p, before=40, after=40)
            r1 = p.add_run(f"{key} : ")
            r1.font.bold = True; r1.font.color.rgb = C_MID
            r2 = p.add_run(val);  r2.font.color.rgb = C_DARK

        doc.add_paragraph()
        self._divider(doc, "38A3E8", thick=10)
        doc.add_page_break()

    # ── section 1 ─────────────────────────────────────────────────────────────

    def _s1_demo(self, doc, metrics: Dict):
        self._h(doc, "1. Données Démographiques & Modes de Vie")

        # 1.1 tableau démographie
        self._h(doc, "1.1 Données démographiques", level=2)
        demo = metrics.get("demographics", {})
        men  = demo.get("men",   {})
        women= demo.get("women", {})

        tbl = doc.add_table(rows=2, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        headers = ["Hommes", "Femmes", "Âge moyen"]
        values  = [
            f"{men.get('pct', 0):.1f}%  (n={men.get('n', 0)})",
            f"{women.get('pct', 0):.1f}%  (n={women.get('n', 0)})",
            f"{demo.get('avg_age', 0)} ans",
        ]
        for i, (h, v) in enumerate(zip(headers, values)):
            ch = tbl.rows[0].cells[i]; cv = tbl.rows[1].cells[i]
            ch.width = Inches(2.0);     cv.width = Inches(2.0)
            _cell_bg(ch, "0F2340");     _cell_borders(ch, "0F2340")
            _cell_bg(cv, "EDF5FD");     _cell_borders(cv, "D6E8F7")
            ph = ch.paragraphs[0]; ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rh = ph.add_run(h)
            rh.font.bold = True; rh.font.color.rgb = C_WHITE; rh.font.size = Pt(10)
            pv = cv.paragraphs[0]; pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rv = pv.add_run(v)
            rv.font.bold = True; rv.font.color.rgb = C_DARK; rv.font.size = Pt(11)

        doc.add_paragraph()

        # 1.2 modes de vie
        self._h(doc, "1.2 Indicateurs de modes de vie", level=2)
        ls = metrics.get("lifestyle", {})
        ls_map = {
            "tobacco":     ("Tabagisme",                  30, True),
            "alcohol":     ("Consommation d'alcool",      30, True),
            "sedentarity": ("Sédentarité",                35, True),
            "chronic":     ("Maladie chronique",          20, True),
            "overweight":  ("Surpoids & Obésité",         30, True),
        }
        found = False
        for k, (lbl, thr, inv) in ls_map.items():
            if k in ls:
                self._kpi(doc, lbl, ls[k]["pct"], ls[k]["n"], good=thr, inv=inv)
                found = True
        if not found:
            self._note(doc, "Aucun indicateur de mode de vie détecté dans ce fichier.")
        doc.add_paragraph()

    # ── section 2 ─────────────────────────────────────────────────────────────

    def _s2_results(self, doc, metrics: Dict):
        self._h(doc, "2. Résultats Principaux")

        # 2.1 seuils
        self._h(doc, "2.1 Méthodologie (Seuils Théoriques)", level=2)
        p = doc.add_paragraph(style="Normal"); _spacing(p, after=100)
        p.add_run("Ce rapport utilise les ").font.color.rgb = C_DARK
        rb = p.add_run("seuils théoriques")
        rb.font.bold = True; rb.font.color.rgb = C_BLUE
        p.add_run(
            " basés sur le point médian de l'échelle Likert (1–4). "
            "Ces indicateurs sont conventionnels et non cliniques."
        ).font.color.rgb = C_DARK

        key_scores = [
            "Dem_score","Lat_score","SS_score",
            "comp_score","auto_score","sup_score","col_score",
            "rec_score","equ_score","cult_score","sat_score",
            "adq_resources_score","adq_role_score",
        ]
        rows_data = [(k, SCORE_LABELS.get(k, k))
                     for k in key_scores if k in THRESHOLDS]

        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for cell, txt in zip(tbl.rows[0].cells, ["Dimension", "Seuil théorique"]):
            _cell_bg(cell, "0F2340"); _cell_borders(cell, "0F2340")
            ph = cell.paragraphs[0]; ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ph.add_run(txt)
            r.font.bold = True; r.font.color.rgb = C_WHITE; r.font.size = Pt(10)
        for i, (k, lbl) in enumerate(rows_data):
            row = tbl.add_row()
            bg = "F5F9FF" if i % 2 == 0 else "FFFFFF"
            _cell_bg(row.cells[0], bg); _cell_borders(row.cells[0], "D6E8F7")
            _cell_bg(row.cells[1], bg); _cell_borders(row.cells[1], "D6E8F7")
            row.cells[0].paragraphs[0].add_run(lbl).font.color.rgb = C_DARK
            pv = row.cells[1].paragraphs[0]; pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rv = pv.add_run(str(THRESHOLDS[k]))
            rv.font.bold = True; rv.font.color.rgb = C_BLUE

        doc.add_paragraph()

        # 2.2 quadrants — tableau 4 cases colorées
        self._h(doc, "2.2 Quadrants de Karasek", level=2)
        quadrants = metrics.get("quadrants", {})
        quad_defs = [
            ("Tendu",   "Demande élevée, contrôle faible"),
            ("Actif",   "Demande et contrôle élevés"),
            ("Passif",  "Demande et contrôle faibles"),
            ("Détendu", "Contrôle élevé, demande modérée"),
        ]
        tbl_q = doc.add_table(rows=3, cols=4)
        tbl_q.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_q.style = "Table Grid"
        for ci, (qname, qsub) in enumerate(quad_defs):
            data = quadrants.get(qname, {})
            pct  = data.get("pct",   0)
            cnt  = data.get("count", 0)
            hx   = QUAD_HEX.get(qname, "94A3B8")

            # Row 0 — nom
            c0 = tbl_q.rows[0].cells[ci]
            _cell_bg(c0, hx); _cell_borders(c0, hx)
            p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(qname.upper())
            r0.font.bold = True; r0.font.color.rgb = C_WHITE; r0.font.size = Pt(10)

            # Row 1 — %
            c1 = tbl_q.rows[1].cells[ci]
            _cell_bg(c1, "F5F9FF"); _cell_borders(c1, "D6E8F7")
            p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(f"{pct:.1f}%")
            r1.font.bold = True; r1.font.size = Pt(16)
            r1.font.color.rgb = _hex_rgb(hx)

            # Row 2 — n + sous-label
            c2 = tbl_q.rows[2].cells[ci]
            _cell_bg(c2, "FFFFFF"); _cell_borders(c2, "D6E8F7")
            p2 = c2.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(f"n = {cnt}\n").font.color.rgb = C_MID
            r2b = p2.add_run(qsub)
            r2b.font.size = Pt(7); r2b.font.italic = True; r2b.font.color.rgb = C_GRAY

        doc.add_paragraph()

        # 2.3 RPS
        self._h(doc, "2.3 Prévalence des RPS", level=2)
        strain = metrics.get("strain_prevalence", {})
        if "Job_strain" in strain:
            self._kpi(doc, "Job Strain — Tension au travail",
                      strain["Job_strain"]["pct"], strain["Job_strain"]["n"],
                      good=25, inv=True)
        if "Iso_strain" in strain:
            self._kpi(doc, "Iso-Strain — Isolement au travail",
                      strain["Iso_strain"]["pct"], strain["Iso_strain"]["n"],
                      good=25, inv=True)
        self._note(doc, "Seuil de vigilance indicatif : > 25% (non clinique).")
        doc.add_paragraph()

        # 2.4 stress
        self._h(doc, "2.4 Indicateurs clés de stress au travail", level=2)
        stress = metrics.get("stress_indicators", {})
        for key, lbl, thr, inv in [
            ("autonomy",       "Autonomie décisionnelle (% élevé)", 50, False),
            ("workload",       "Charge mentale perçue (% élevé)",   50, True),
            ("social_support", "Soutien social (% élevé)",          50, False),
        ]:
            if key in stress:
                self._kpi(doc, lbl, stress[key]["pct"], stress[key]["n"],
                          good=thr, inv=inv)
        doc.add_paragraph()

        # 2.5 RH — tableau coloré
        self._h(doc, "2.5 Climat Organisationnel (Scores RH)", level=2)
        rh = metrics.get("rh_scores", {})
        if rh:
            tbl_rh = doc.add_table(rows=1, cols=3)
            tbl_rh.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_rh.style = "Table Grid"
            for cell, txt in zip(tbl_rh.rows[0].cells, ["Dimension", "% Élevé", "n"]):
                _cell_bg(cell, "0F2340"); _cell_borders(cell, "0F2340")
                ph = cell.paragraphs[0]; ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = ph.add_run(txt)
                r.font.bold = True; r.font.color.rgb = C_WHITE; r.font.size = Pt(10)
            for i, (col, data) in enumerate(rh.items()):
                row = tbl_rh.add_row()
                bg  = "F5F9FF" if i % 2 == 0 else "FFFFFF"
                pct_v = data["pct"]
                vc = "22C55E" if pct_v >= 50 else "EF4444"
                _cell_bg(row.cells[0], bg); _cell_borders(row.cells[0], "D6E8F7")
                row.cells[0].paragraphs[0].add_run(data["label"]).font.color.rgb = C_DARK
                _cell_bg(row.cells[1], bg); _cell_borders(row.cells[1], "D6E8F7")
                pp = row.cells[1].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rp = pp.add_run(f"{pct_v:.1f}%")
                rp.font.bold = True; rp.font.color.rgb = _hex_rgb(vc)
                _cell_bg(row.cells[2], bg); _cell_borders(row.cells[2], "D6E8F7")
                pn = row.cells[2].paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pn.add_run(str(data["n"])).font.color.rgb = C_GRAY
        else:
            self._note(doc, "Scores RH non disponibles dans ce fichier.")
        doc.add_paragraph()

    # ── section 3 ─────────────────────────────────────────────────────────────

    def _s3_figures(self, doc, figures: Dict[str, bytes]):
        if not figures:
            return
        self._h(doc, "3. Visualisations")

        fig_meta = [
            ("karasek_scatter",     "3.1 Grille MAPP — Demande × Latitude",         Inches(5.6)),
            ("karasek_heatmap",     "3.2 Répartition Karasek par Direction",         Inches(6.0)),
            ("strain_bars",         "3.3 Prévalence des RPS & Quadrants",            Inches(5.0)),
            ("rh_radar",            "3.4 Radar des dimensions organisationnelles",   Inches(4.6)),
            ("age_pyramid",         "3.5 Pyramide des âges",                         Inches(4.8)),
            ("genre_barplot",       "3.6 Répartition par Genre",                     Inches(4.4)),
            ("tranche_age_barplot", "3.7 Répartition par Tranche d'âge",             Inches(4.4)),
            ("direction_barplot",   "3.8 Répartition par Direction",                 Inches(4.6)),
            ("csp_barplot",         "3.9 Répartition par CSP",                       Inches(4.6)),
        ]
        for key, title, width in fig_meta:
            if key not in figures or not figures[key]:
                continue
            self._h(doc, title, level=2)
            doc.add_picture(io.BytesIO(figures[key]), width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    # ── pied de document ──────────────────────────────────────────────────────

    def _footer(self, doc):
        doc.add_paragraph()
        self._divider(doc, "D6E8F7", thick=4)
        for note in [
            "Les seuils utilisés sont des seuils théoriques (point médian Likert 1–4). "
            "Ils sont conventionnels et ne constituent pas des seuils cliniques diagnostiques.",
            "Job Strain : demande psychologique élevée + latitude décisionnelle faible.",
            "Iso-Strain : Job Strain + faible soutien social.",
            f"Rapport généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Wave-CI Platform.",
        ]:
            self._note(doc, f"• {note}")

    # ── point d'entrée ────────────────────────────────────────────────────────

    def generate(
        self,
        metrics: Dict[str, Any],
        figures: Optional[Dict[str, bytes]] = None,
    ) -> bytes:
        """
        Génère le rapport Word et retourne les bytes .docx.

        Parameters
        ----------
        metrics : dict retourné par KarasekAnalytics.compute()
        figures : dict {key: png_bytes} retourné par KarasekVisualizations.generate_all()
        """
        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.8)
            section.right_margin  = Cm(2.8)

        total = metrics.get("demographics", {}).get("total", "N/A")

        self._cover(doc, total)
        self._s1_demo(doc, metrics)
        self._s2_results(doc, metrics)
        self._s3_figures(doc, figures or {})
        self._footer(doc)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
